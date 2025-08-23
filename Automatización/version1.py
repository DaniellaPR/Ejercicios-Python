#### PROGRAMA DE AUTOMATIZACIÓN DE INFORMES DE CONTRATOS CIVILES DE DTIC ----------------------------------------

""""
Versión N°4
Fecha: 2025/08/21

Descripción:

El presente programa se encarga de automatizar la generación de los presentes informes del personal que se encuentra bajo contratos civiles:

- Informe de Productos
- Informe de Actividades y Productos Entregados
- Informe de Aceptación de los Productos

Toda la información que se inserta en los informes se obtiene de su respectivo TDR. Para la presente prueba se utilizó el TDR: CTDR_Asistente_Desarrollo_Informático

Para poder realizar la automatización tambien se utilizaron una serie de plantillas de los distintos informes, los cuales presentan una notación Jinja2 para poder 
insertar marcadores en donde se deba de reemplazar información cambiante en los reportes.

A continuación el código fuente del programa:

"""


#### LIBRERÍAS UTILIZADAS ----------------------------------------

from docxtpl import DocxTemplate
from docx import Document
from datetime import datetime
import calendar
import pandas as po
import os
import re
import sys
from docx.oxml import OxmlElement
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from pathlib import Path


################### SECCIÓN #1: FUNCIONES #######################


#### FUNCION PARA ENCONTRAR EL TDR EN EL DIRECTORIO ----------------------------------------

BASE_DIR = Path(__file__).resolve().parent # Ruta absoluta para encontrar Excel basada en .py

# Funcion para encontrar el archivo de Excel
def load_excel(filename: str, sheet: str):
    path = BASE_DIR / filename

    # Si se desea visualizar la dirección específica del Script y el TDR.
    """
    print(f"CWD: {Path.cwd()}")
    print(f"Script dir: {BASE_DIR}")
    print(f"Buscando Excel en: {path}")
    """

    if not path.exists():
        print(f"No se encontró: {path.name}")
        print(" XLSX disponibles en la carpeta del script:")
        for p in BASE_DIR.glob("*.xlsx"):
            print(" -", p.name)
        sys.exit(1)  

    try:
        return po.read_excel(path, sheet_name=sheet, engine="openpyxl")
    except Exception as e:
        print(f"Error leyendo '{path.name}': {e}")
        sys.exit(1)

#### FUNCIONES PARA OTORGAR FORMATO A LAS TABLAS GENERADAS POR Docx ----------------------------------------

# Función para añadir un borde visible a las tablas generadas.

def set_borders(table):

    tbl = table._element
    tblBorders = OxmlElement('w:tblBorders')

    for border_name in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        border = OxmlElement(f"w:{border_name}")
        border.set(qn('w:val'), 'single')
        border.set(qn('w:sz'), '8')
        border.set(qn('w:space'), '0')
        border.set(qn('w:color'), '000000')
        tblBorders.append(border)
    tbl.tblPr.append(tblBorders)

# Función para darle formato al texto generado dentro de las tablas generadas.

def aplicar_formatos(celda, negrita = False):

    for parrafo in celda.paragraphs:
        for run in parrafo.runs:
            run.font.name = "Century Gothic"
            run.font.size = Pt(11)
            run.bold = negrita


#### FUNCIONES PARA OBTENER LOS PLAZOS DE ENTREGA DE LOS PRODUCTOS SEGÚN LA INFORMACIÓN DEL TDR ----------------------------------------


# Funcion para leer los plazos de la metodología y juntar la información en strings reutilizables.

def fucionar_por_plazos (items, marcador = "*", espacio_despues=True, unir_inicio = False):

    resultado = []
    acumulado = False
    buffer = []
    inicio = []

    for token in items:
        if token == marcador:

            if unir_inicio and inicio:
                resultado.append(" ".join(inicio))
                inicio = []
            elif inicio:
                resultado.extend(inicio)
                inicio = []

            if acumulado:
                prefijo = (marcador + " ") if espacio_despues else marcador
                resultado.append(prefijo + " ".join(buffer))
                buffer = []
            acumulado = True

        else:
            if acumulado:
                buffer.append(token)
            else:
                resultado.append(token)
    if acumulado:
        prefijo = (marcador + " ") if espacio_despues else marcador
        resultado.append(prefijo + " ".join(buffer))
    elif inicio:
        if unir_inicio:
            resultado.append("".join(inicio))
        else:
            resultado.extend(inicio)
    
    return resultado


# Función para poder separar del string obtenido de la sección de Metodología del TDR para así obtener una lista con los productos junto a sus plazos respectivos
def separar_periodo (periodo: str):
    a, b = re.split(r"\s*hasta\s*", periodo, flags=re.IGNORECASE)
    return a.strip(), b.strip()


# Función para identificar si un Producto se encuentra en dos plazos distintos según lo indicado en la sección de Metodología del TDR. 
# En el caso de que se encuentre en dos plazos, la fecha de inicio del producto será la primera encontrada y la fecha de finalización 
# será la última fecha mencionada del producto en la sección de Metodología.

def plazos_unidos(lista_1,lista_2):
    titulos1, periodo1 = lista_1[:-1], lista_1[-1]
    titulos2, periodo2 = lista_2[:-1], lista_2[-1]

    ini1, fin1 = separar_periodo(periodo1)
    ini2, fin2 = separar_periodo(periodo2)

    orden = titulos2 + [t for t in titulos1 if t not in titulos2]

    plazos_por_titulo = {}

    for t in orden:
        if t in titulos1 and t in titulos2:
            plazos_por_titulo[t] = f"{ini1} hasta {fin2}"
        elif t in titulos1:
            plazos_por_titulo[t] = periodo1
        else:
            plazos_por_titulo[t] = periodo2
    return [plazos_por_titulo[t] for t in orden]


#### FUNCIONES PARA LA CREACIÓN DE CARPETAS Y DOCUMENTOS GUARDADOS ----------------------------------------

# Se extraen valores de fechas utilizando la fecha actual como dato para calcular el inicio del periodo y el final del periodo.
fecha_actual = datetime.now()
inicio_periodo = fecha_actual.replace(day=1).strftime("%d-%m-%Y")
ultimo_dia = calendar.monthrange(fecha_actual.year, fecha_actual.month)[1]
fin_periodo = fecha_actual.replace(day=ultimo_dia).strftime("%d-%m-%Y")
meses = {
    1: "enero", 2:"febrero", 3:"marzo", 4:"abril", 5:"mayo", 6:"junio", 7:"julio", 8:"agosto", 9:"septiembre", 10:"octubre", 11:"noviembre", 12:"diciembre"
}
mes_Actual=meses[fecha_actual.month]
año_Actual = fecha_actual.year


# Función para detectar el escritorio en la computadora.
def get_desktop() -> Path:
    candidatos = [
        Path(os.environ.get("USERPROFILE", "")) / "Desktop",
        Path(os.environ.get("OneDrive", "")) / "Desktop",
        Path(os.environ.get("ONEDRIVE", "")) / "Desktop",
        Path.home() / "Desktop",
    ]
    for p in candidatos:
        if p and p.exists():
            return p
        
    # Si no se encuentra ninguno de los anteriores se guarda en la carpeta del usuario

    return Path.home()

# Función para crear carpetas en donde almacenar los informes.
def ensure_dir(path: Path) -> Path:
    path.mkdir(parents=True, exist_ok=True)
    return path


# Se procede a abrir el archivo de excel en donde se encuentra el TDR haciendo uso de las funciones codificadas en la sección anterior.
try:
    df = load_excel("CTDR_Asistente_Desarrollo_Informatico.xlsx", "TDR Asistente v02")
except FileNotFoundError:
    print("No se encontró el archivo CTDR_Asistente_Desarrollo_Informatico.xlsx")
except Exception as e:
    print("Error al abrir el archivo CTDR_Asistente_Desarrollo_Informatico.xlsx: "+ str(e))

# Se procede a abrir el archivo de excel en donde se encuentra la información de los funcionarios y sus honorarios

try:
    df_h = load_excel("Matriz de honorarios y valores a pagar DITIC.xlsx", "Hoja1")
except FileNotFoundError:
    print("No se encontró el archivo Matriz de honorarios y valores a pagar DITIC.xlsx")
except Exception as e:
    print("Error al abrir el archivo Matriz de honorarios y valores a pagar DITIC.xlsx" + str(e))

#Variables temporales para el nombre de carpeta general mensual y por puesto
mes_Carpeta = f"{mes_Actual}-{año_Actual}"
# Función para guardar el .docx
def save_docx(doc, file_name: str, folder_name: str = "Informe_General", puesto: str = "") -> Path:
    invalid = '<>:"/\\|?*' # Se limpian los caracteres inválidos.
    safe_file = "".join(c for c in file_name if c not in invalid).strip() or "informe.docx"
    safe_folder = "".join(c for c in folder_name if c not in invalid).strip() or "Carpeta_Funcionario"

    #dest_dir = ensure_dir(get_desktop() / mes_Carpeta / puesto / safe_folder)
    dest_dir = ensure_dir(get_desktop() / mes_Carpeta/puesto.strip()/safe_folder)
    dest_path = dest_dir / safe_file
    doc.save(str(dest_path))  
    print(f"Documento guardado en: {dest_path}")
    return dest_path

def limpio (x: object) -> str:

    if po.isna(x):
        return ""
    s = str(x).strip()
    s = re.sub(r"\s+", " ",s)
    return s

################### SECCIÓN #2: CARGA DE TDR Y OBTENCIÓN DE DATOS ##########################
# Para poder constatar la correcta carga de la Matriz de Honorarios y Valores a Pagar, se genera una lista con la información de los distintos funcionarios




funcionario_inicial = 3
columna_nombres = 1
columna_apellidos = 2

lista_funcionarios = []

r = funcionario_inicial

while r < len(df_h):
    nombres = limpio(df_h.iat[r,columna_nombres])
    apellidos = limpio(df_h.iat[r,columna_apellidos])

    if not nombres and not apellidos:
        break
    completo = " ".join(p for p in (nombres, apellidos) if p)
    lista_funcionarios.append(completo)

    r+=1



columna_cargo = 5
lista_cargo_func = []

r_2 = funcionario_inicial

while r_2 < len(df_h):
    v_2 = df_h.iat[r_2, columna_cargo]
    if po.isna(v_2) or (isinstance(v_2, str) and v_2.strip() == ""):
        break
    lista_cargo_func.append(v_2)
    r_2 +=1

columna_cedulas = 4
lista_cedula_func = []

r_3 = funcionario_inicial

while r_3 < len(df_h):
    v_3 = df_h.iat[r_3, columna_cedulas]
    if po.isna(v_3) or (isinstance(v_3, str) and v_3.strip() == ""):
        break
    lista_cedula_func.append(v_3)
    r_3 +=1

columna_honorarios = 6
lista_honorario_func = []

r_4 = funcionario_inicial

while r_4 < len(df_h):
    v_4 = df_h.iat[r_4, columna_honorarios]
    if po.isna(v_4) or (isinstance(v_4, str) and v_4.strip() == ""):
        break
    lista_honorario_func.append(v_4)
    r_4 +=1




# Para constatar la correcta carga del TDR, se mantiene un registro de los productos y actividades encontradas en el TDR.
try:
    numero_inicial_prod = df.iloc[33:,0]
    numeros = numero_inicial_prod[numero_inicial_prod.apply(lambda x: isinstance(x, (int)))]
    ultimo_valor = numeros.iloc[-1] if not numeros.empty else None
except Exception:
    ultimo_valor=0 # Con el último valor es posible recorrer exactamente las actividades y productos para almacenarlos en listas

#Lista de Actividades y productos
actividades_lista, productos_lista = [], []

# Posición Inicial Actividades
columna_actividades = 1
fila_actividades = 33 # Recordemos que usamos ultimo_valor para saber cuantas posiciones recorrer

# Posicion Inicial de los Productos
columna_productos = 6
fila_productos = 33 # Recordemos que usamos ultimo_valor para saber cuantas posiciones recorrer

# Para poder utilizar la información sobre las actividades y productos, se procede a extraer su información del TDR y se la almacena en listas.
if ultimo_valor>0:
    try:
        # Extraen valores de actividad
        serie_actividades = df.iloc[fila_actividades:fila_actividades + ultimo_valor, columna_actividades]
        actividades_lista = serie_actividades.tolist()
    except Exception:
        print("No se pudo extraer los datos de Actividades")
    try:
        # Extraen valores de los Productos
        serie_productos = df.iloc[fila_productos: fila_productos + ultimo_valor, columna_productos]
        productos_lista = serie_productos.tolist()
    except Exception:
        print("No se pudo extraer los datos de Productos")

# Una vez creadas las listas, se extrae información importante para el contexto generalizado
# Se extraen los valores del proyecto, puesto y honorarios
try:
    proyecto = df.iloc[9,4].upper()
    puesto = df.iloc[10,4]
    honorario = df.iloc[68,4]
except Exception:
    proyecto, puesto, honorario="No disponible", "No disponible", 0




# Se extrae de la metodología los plazos de entrega de cada producto de manera individual.
metodologia = df.iloc[23,0]
lst_metodologia = metodologia.split()
indices = [i for i, palabra in enumerate(lst_metodologia)
            if palabra ==("Fecha:") ]  
plazos = []
for i in indices:
    rango = lst_metodologia[i+2:i+10]
    plazos.append(" ".join(rango))
if len(indices)>2:
    segunda = lst_metodologia[indices[1]+2]
    print(" ".join(segunda))
try:
    plazo_1=plazos[0]
    plazo_2=plazos[1]
except Exception:
    print("Error al cargar los plazos")


# Obtención de la seccion de fechas para el periodo escrito en texto natural.
hoy = datetime.now()
inicio = hoy.replace(day=1)
ultimo = calendar.monthrange(hoy.year, hoy.month)[1]
fin = hoy.replace(day = ultimo)

periodo_incluir_info = f"{inicio.day:02d} al {fin.day:02d} de {meses[hoy.month]} de {hoy.year}"

# Como un paso fundamental, se tiene que realizar la creación de un contexto el cual almacene la ifnormación general a colocar en la maoría de informes.
contexto_plantilla = {

    "proyecto": proyecto,
    "mes": mes_Actual.upper(),

    # Datos quemados ya que no se posee mayor información
    #"funcionario": "BRYAN BENJAMÍN SARABINO CUICHÁN", 
    #"cedula": "172231964-5", 

    #"puesto": puesto,
    #"honorario":honorario,
    "fecha_inicio":inicio_periodo,
    "fecha_fin":fin_periodo,
    "periodo":periodo_incluir_info,

    # Información que variará dependiendo de la realización del producto.
    "conclusion_1":"Escribir conclusión #1", 
    "conclusion_2":"Escribir conclusión #2",
    "conclusion_3":"Escribir conclusión #3",
    "recomendacion_1":"Escribir recomendación #1",
    "recomendacion_2":"Escribir recomendación #2",
    "recomendacion_3":"Escribir recomendación #3"

}

################### SECCIÓN #3: MENÚ PRINCIPAL Y EJECUCIÓN DEL PROGRAMA #############################


#### Parte #1: Carga de Plantillas ----------------------------------------

# Para facilitar la obtención de las diferentes plantillas para los informes, se ha creado un Excel donde se almacena, 
# por cada informe, un path con la locaclización exacta de la plantilla a utilizar
try:
    df_pl = load_excel("PLANTILLAS_INFORMES.xlsx", "Hoja1")
except FileNotFoundError:
    print("El arichivo no PLANTILLAS_INFORMES.xlsx no se encontró.")
except Exception as e:
    print("Error al abrir el archivo PLANTILLAS_INFORMES.xlsx.")


#### Parte #2: Menú Principal ----------------------------------------

# Se genera un menú principal el cual brinda la opción de generar los diferentes tipos de informes.
bandera = "s"
while bandera == "s":
    # ---------------- Menú de selección de funcionario -------------------"")
    print("\nPor favor sleccione el funcionario perteneciente a DITIC del cual generará los correspondientes informes: ")
    print("Funcionarios:")
    for i in range(len(lista_funcionarios)):
        print(f"{i+1}. {lista_funcionarios[i]}")
    print("0. Salir")

    funcionario_eleccion_indice = input("Ingrese su selección: ")
    if funcionario_eleccion_indice == "0":
        print("Saliendo del programa...")
        break

    funcionario_seleccionado = lista_funcionarios[int(funcionario_eleccion_indice) - 1]
    cargo_seleccionado = lista_cargo_func[int(funcionario_eleccion_indice) -1]
    cedula_seleccionada = lista_cedula_func[int(funcionario_eleccion_indice) -1]
    honorario_seleccionado = lista_honorario_func[int(funcionario_eleccion_indice) -1]

    print(funcionario_seleccionado)
    print(cargo_seleccionado)
    print(honorario_seleccionado)
    print(f"Usted a seleccionado a: {lista_funcionarios[int(funcionario_eleccion_indice) - 1]}")

    salir_programa = False
    bandera2 = "s"
    while bandera2 == "s":
        print(" -------------------- Menú principal -------------------")
        print("1. Informe de Productos")
        print("2. Informe de Actividades y Productos Entregados")
        print("3. Informe de Aceptación de los Productos Recibidos a Satisfacción")
        print("4. Elegir otro funcionario")
        print("5. Salir del programa")

        seleccion_menu = input("Indique la opción (1 - 5): ")

        bandera3 = "s"
        while bandera3 == "s":
            if seleccion_menu in ["1", "2", "3", "4", "5"]:
                    bandera3 = "n"
            else:
                print("Opción inválida, por favor seleccione una opción válida.")
                seleccion_menu = input("Indique el informe (1 - 5): ")
        
            
        ##### Menú Pt1. Generación de Informe de Productos ------------------------------------------------------
        if (seleccion_menu == "1"):
            if ultimo_valor==0:
                print("No hay productos para generar")
            else:
                print("\nElija cual informe realizar")
                for i in range(ultimo_valor):
                    print(f"{i + 1}. Producto {i + 1}")
                print(f"{ultimo_valor + 1}. Generar todos los Informes")
                print(f"{ultimo_valor + 2}. Volver al menú principal")

                usuario_seleccion = input(f"Digite su elección (1 - {ultimo_valor+1}): ")
                if not usuario_seleccion.isdigit():
                    print("Debe ingresar un número")
                else:
                    usuario_seleccion = int(usuario_seleccion)

                    # En el caso de que el usuario solamente quiera generar el ifnorme de un producto en específico se realiza lo siguiente:
                    if 1<= usuario_seleccion <= ultimo_valor:
                        try:
                            path_informe = df_pl.iloc[0,2] # Se selcciona la celda con el path a la plantilla del Informe de Productos.
                            plantilla_prod = DocxTemplate(path_informe)

                            # Se utiliza un contexto personalizado utilizando al contexto base junto con la actividad y producto seleccionados.
                            contexto_actividad_prod = {
                                **contexto_plantilla,
                                "funcionario":funcionario_seleccionado,
                                "cedula": cedula_seleccionada,
                                "puesto": cargo_seleccionado,
                                "producto":productos_lista[int(usuario_seleccion) - 1],
                                "actividad":actividades_lista[int(usuario_seleccion) - 1],
                                "numero": usuario_seleccion
                            }

                            # Se renderiza el contexto y se guarda el informe creando una carpeta y un documento word mediante las funciones de la primera sección.
                            plantilla_prod.render(contexto_actividad_prod)
                            
                            #Confirmación de guardado del informe
                            confirmacion = input(f"Seguro desea guardar el informe? (s/n): ")
                            if confirmacion.lower() == 's':
                                save_docx(plantilla_prod,f"Informe_de_Producto_{usuario_seleccion}.docx", funcionario_seleccionado, cargo_seleccionado)
                            else:
                                print("Informe no guardado.")


                        except Exception as e:
                            print(f"Error al generar el informe: {e}")
                        print()
                        print()

                    # En el caso de que el usuario eliga crear todos los informes de todos los productos de manera automática, se realiza lo siguiente:
                    elif usuario_seleccion == ultimo_valor+1:
                        for i in range(1, ultimo_valor+1):
                            try:
                                path_informe = df_pl.iloc[0,2] # Se selcciona la celda con el path a la plantilla del Informe de Productos.
                                plantilla_prod = DocxTemplate(path_informe)

                                # Se utiliza un contexto personalizado utilizando al contexto base junto con la lista complea de actividades y productos del TDR.
                                contexto_actividad_prod = {
                                    **contexto_plantilla,
                                    "funcionario":funcionario_seleccionado,
                                    "cedula": cedula_seleccionada,
                                    "puesto": cargo_seleccionado,
                                    "producto":productos_lista[i - 1],
                                    "actividad":actividades_lista[i - 1],
                                    "numero": i
                                }

                                # Se renderiza el contexto y se guarda el informe creando una carpeta y un documento word mediante las funciones de la primera sección.
                                plantilla_prod.render(contexto_actividad_prod)
                                
                                #Confirmación de guardado del informe
                                confirmacion = input(f"Seguro desea guardar el informe? (s/n): ")
                                if confirmacion.lower() == 's':
                                    save_docx(plantilla_prod, f"Informe_de_Producto_{i}.docx", funcionario_seleccionado, cargo_seleccionado)
                                else:
                                    print("Informe no guardado.")
                                
                            except Exception as e:
                                print(f"Error en el informe {i}: {e}")
                        print()
                        print()
                    elif usuario_seleccion == ultimo_valor + 2:
                            print("Volviendo al menú principal...")
                            bandera2 = "s"
                    else:
                        if usuario_seleccion > ultimo_valor + 2:
                            print("Error fuera de rango.")
                            print()

        ##### Menú Pt2. Generación de Informe de Actividades y Productos Entregados ------------------------------------------------------
        elif (seleccion_menu == "2"):
            print()
            print()
            print("Indique cuantos productos y/o actividades tiene el informe:")
            for i in range(ultimo_valor):
                print(f"{i + 1}. Producto / Actividad {i + 1}")
            print(f"{ultimo_valor + 1}. Volver al menú principal")
            print()

            print("Digite la letra 'T' para generar un informe con todos los Productos / Actividades realizadas.")
            print()
            print("Indique de que Producto / Actividad realizar el reporte.")
            print("Recuerde que puede digitar más de una opción, separelo por espacios.")
            usuario_seleccion = input(f"Ingrese sus selecciones (1 - {ultimo_valor+1} / T): ")

            # Se le habilita al usuario la opción de generar un informe que contenga más de un producto, ya que podrá digitar el número del producto separado por espacios.
            if (usuario_seleccion != 'T' and usuario_seleccion != str(ultimo_valor + 1)):
                lista_seleccion = usuario_seleccion.split()

                # Se controla que las opciones seleccionadas no superen al último valor
                lista_depurada = []
                for i in lista_seleccion:
                    if ((int(i) >= ultimo_valor + 2)):
                        print("Lo sentimos, a ingresado un valor invalido.")
                        print(f"No existe actualmente un Producto / Actividad {i}")
                        print("Se removerá a esta opción de su selección.")
                    else:
                        lista_depurada.append(int(i))

                # En las presentes listas se almacenan los productos y actividades seleccionados.
                productos_seleccionados = []
                actividades_seleccionadas = []

                for i in lista_depurada:
                    productos_seleccionados.append(productos_lista[i - 1])
                    actividades_seleccionadas.append(actividades_lista[i -1])
                print()
                print("A continuación se generará su informe")

                df_pl = load_excel("PLANTILLAS_INFORMES.xlsx", "Hoja1") 
        
                path_informe = df_pl.iloc[1,2] # Se carga la plantilla del informe accediendo a su celda en el Excel.

                plantilla_prod_act = DocxTemplate(path_informe)

                # Se utiliza un contexto personalizado que permite utilizar solamente los producos y actividades seleccionados.
                contexto_actividad_prod = {
                    **contexto_plantilla,
                    "funcionario":funcionario_seleccionado,
                    "cedula": cedula_seleccionada,
                    "puesto": cargo_seleccionado,
                    "productos_lista":productos_seleccionados,
                    "actividades_lista":actividades_seleccionadas,
                }

                # Se renderiza el documento y se lo almacena en su respectiva carpeta.
                plantilla_prod_act.render(contexto_actividad_prod)
                
                
                #Confirmación de guardado del informe
                confirmacion = input(f"Seguro desea guardar el informe? (s/n): ")
                if confirmacion.lower() == 's':
                    save_docx(plantilla_prod_act, f"Informe_de_Actividad_Producto_Realizados_Total.docx", funcionario_seleccionado, cargo_seleccionado)
                else:
                    print("Informe no guardado.")
                
                print()
                print()

            elif (usuario_seleccion == str((ultimo_valor + 1))):
                            print("Volviendo al menú principal...")
                            bandera2 = "s"
            # En el caso de que el usuario quiera generar un informe con todos los prodcutos y actividades del TDR se realiza lo siguiente:
            else:
                print()
                df_pl = load_excel("PLANTILLAS_INFORMES.xlsx", "Hoja1")
        
                path_informe = df_pl.iloc[1,2]

                plantilla_prod_act = DocxTemplate(path_informe)

                contexto_actividad_prod = {
                    **contexto_plantilla,
                    "funcionario":funcionario_seleccionado,
                    "cedula": cedula_seleccionada,
                    "puesto": cargo_seleccionado,
                    "productos_lista":productos_lista,
                    "actividades_lista":actividades_lista,
                }
                plantilla_prod_act.render(contexto_actividad_prod)
                
                #Confirmación de guardado del informe
                confirmacion = input(f"Seguro desea guardar el informe? (s/n): ")
                if confirmacion.lower() == 's':
                    save_docx(plantilla_prod_act, f"Informe_de_Actividad_Producto_Realizados_Total.docx", funcionario_seleccionado, cargo_seleccionado)
                else:
                    print("Informe no guardado.")
                
                print()
                print()
        ##### Menú Pt3. Generación de Informe de Aceptación de Productos ------------------------------------------------------
        # Esta opción se encuentra en desarrollo, en la presente versión, se genera el informe con todos los productos y actividades del TDR

        elif (seleccion_menu == "3"):

            plazos_de_informe = []

            print()
            usuario_seleccion = input("Digite la letra 'T' para generar un informe con todos los Productos / Actividades realizadas.\nPresione 0 para volver al menú principal.\nIngrese su selección: ")
            

            if (usuario_seleccion == 'T'):

                # Considerando que existen plazos para productos, según la sección de metodología, los cuales difieren por producto, 
                # se procede a extraer la información precisa de la metodología para poder identificar los plazos exactos de cada 
                # producto haciendo uso de las funciones de la primera sección.
                
                listas_para_unir = []
                metodologia = df.iloc[23,0]
                lst_metodologia = metodologia.split()
                indice_1 = [i for i, palabra in enumerate(lst_metodologia)
                        if palabra ==("Fecha:") ]  # Utilizar solo el primer elemento
                
                indice_2 = [i for i, palabra in enumerate(lst_metodologia)
                            if palabra == ("Honorarios:")] # Utilizar solo el primer elemento

                lista_plazos = lst_metodologia[indice_1[0]:indice_2[0]]
                
                plazos_separados = []
                actual = []
                for elemento in lista_plazos:
                    if elemento == "Fecha:":

                        if actual:
                            plazos_separados.append(actual)
                        actual = [elemento]
                    else:
                        actual.append(elemento)
                if actual:
                    plazos_separados.append(actual)
                for plazos in plazos_separados:
                    plazos_semi_procesados = fucionar_por_plazos(plazos, marcador = "*",espacio_despues=True ,unir_inicio=True)
                    
                    informes_plazo = plazos_semi_procesados[10:].copy()
                    plazos_texto = ""

                    for i in range(2, 10):
                        plazos_texto += plazos_semi_procesados[i] + " "
                    
                    plazos_texto = plazos_texto.strip()
                    informes_plazo.append(plazos_texto)
                    listas_para_unir.append(informes_plazo)

                if len(listas_para_unir) >= 2:
                    lista1, lista2 = listas_para_unir[0], listas_para_unir[1]
                    plazos_final = plazos_unidos(lista1, lista2)
                else:
                    plazos_final = listas_para_unir[0] if listas_para_unir else []

                # Se procede a crear la tabla de la sección de Objetivos del Informe
                # Las tablas generadas, en la presente versión son generadas en archivos word aparte y luego se copian al informe como subdoc.

                doc_1 = Document()

                tabla = doc_1.add_table(rows=3, cols=3)

                
                headers = ["Nombre", "Productos que deben ser entregados por el personal contratado", "Plazo de entrega de los productos"]

                for i, texto in enumerate(headers):
                    celda = tabla.rows[0].cells[i]
                    celda.text = texto
                    aplicar_formatos(celda, negrita = True)

                top_cell = tabla.cell(1,0)
                bottom_cell = tabla.cell(2,0)
                top_cell.merge(bottom_cell)
                top_cell.text = funcionario_seleccionado
                top_cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                aplicar_formatos(top_cell)

                tabla.cell(1,1).text = productos_lista[0]
                tabla.cell(1,2).text = plazos_final[0]
                tabla.cell(2,1).text = productos_lista[1]
                tabla.cell(2,2).text = plazos_final[1]

                for i in range(1,3):
                    for j in range(1,3):
                        aplicar_formatos(tabla.cell(i,j))

                set_borders(tabla)
                doc_1.save("tabla_base.docx")

                df_pl = load_excel("PLANTILLAS_INFORMES.xlsx", "Hoja1")
        
                path_informe = df_pl.iloc[2,2]

                contexto_actividad_prod = {
                    **contexto_plantilla,
                    "funcionario":funcionario_seleccionado,
                    "cedula": cedula_seleccionada,
                    "puesto": cargo_seleccionado,
                    "honorario": honorario_seleccionado,
                    "productos_lista":productos_lista,
                    "actividades_lista":actividades_lista,
                }
                plantilla_prod_acpt = DocxTemplate(path_informe)
                subdoc_1 = plantilla_prod_acpt.new_subdoc("tabla_base.docx")


                # De igual manera se crea la segunda tabla presente en la sección de Productos Recibidos y Fechas de Recepción
                doc_2 = Document()
                tabla2 = doc_2.add_table(rows=3, cols=3)

                headers2 = ["Nombre", "Productos entregados", "Presenta"]

                for i, texto in enumerate(headers2):
                    celda = tabla2.rows[0].cells[i]
                    celda.text = texto
                    aplicar_formatos(celda, negrita=True)

                top_cell2 = tabla2.cell(1, 0)
                bottom_cell2 = tabla2.cell(2, 0)
                top_cell2.merge(bottom_cell2)
                top_cell2.text = funcionario_seleccionado
                top_cell2.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                aplicar_formatos(top_cell2)

                tabla2.cell(1, 1).text = productos_lista[0]
                tabla2.cell(1, 2).text = "Sí"
                tabla2.cell(2, 1).text = productos_lista[1]
                tabla2.cell(2, 2).text = "Sí"

                for i in range(1, 3):
                    for j in range(1, 3):
                        aplicar_formatos(tabla2.cell(i, j))

                set_borders(tabla2)

                doc_2.save("tablas_base.docx")

                df_pl = load_excel("PLANTILLAS_INFORMES.xlsx", "Hoja1")
                path_informe = df_pl.iloc[2, 2]
                plantilla_prod_acpt = DocxTemplate(path_informe)
                subdoc_2 = plantilla_prod_acpt.new_subdoc("tablas_base.docx")
                contexto = {**contexto_plantilla,"tabla_1":subdoc_1, "funcionario":funcionario_seleccionado, "cedula": cedula_seleccionada, "puesto": cargo_seleccionado, "honorario": honorario_seleccionado ,"tabla_2": subdoc_2}
                plantilla_prod_acpt.render(contexto)
                
                #Confirmación de guardado del informe
                confirmacion = input(f"Seguro desea guardar el informe? (s/n): ")
                if confirmacion.lower() == 's':
                    save_docx(plantilla_prod_acpt, "Informe_Aceptacion_Productos.docx", funcionario_seleccionado, cargo_seleccionado)
                else:
                    print("Informe no guardado.")
        
        else:
            if seleccion_menu == "4":
                print("Regresando a la selección de funcionario...")
                print()
                bandera2 = "n"
            elif seleccion_menu == "5":
                print("Saliendo...")
                salir_programa = True
                bandera2 = "n"
                bandera = "n"
            else:
                print("Opción inválida, por favor seleccione una opción válida.")
                print()

    # Finalmente se pregunta al usuario si desea generar otro reporte.
    
    if seleccion_menu != "4" and salir_programa==False:
        bandera = input("¿Desea generar otro informe? (s/n): ").lower()

print()
print("¡Muchas gracias por usar el programa!")
